import os
from datetime import datetime

def saveJD(company, role, jobDescription):
    safe_company = ''.join(c if c.isalnum() else '_' for c in company)
    safe_role = ''.join(c if c.isalnum() else '_' for c in role)
    folder_path = os.path.join('S:\\', 'applications', safe_company, safe_role)
    os.makedirs(folder_path, exist_ok=True)
    now = datetime.now()
    mdd = f"{now.month:02}{now.day:02}"
    jd_file_path = os.path.join(folder_path, f"jd{mdd}.txt")
    with open(jd_file_path, "w", encoding="utf-8") as f:
        f.write(jobDescription)
    return jd_file_path

def saveResumeDocx(company, role, resumeMarkdown):
    safe_company = ''.join(c if c.isalnum() else '_' for c in company)
    safe_role = ''.join(c if c.isalnum() else '_' for c in role)
    folder_path = os.path.join('S:\\', 'applications', safe_company, safe_role)
    os.makedirs(folder_path, exist_ok=True)
    now = datetime.now()
    mdd = f"{now.month:02}{now.day:02}"
    resume_file_name = f"resume{mdd}.docx"
    resume_file_path = os.path.join(folder_path, resume_file_name)


    # Remove code block markers if present
    cleaned = resumeMarkdown.strip()
    if cleaned.startswith('```markdown'):
        cleaned = cleaned[len('```markdown'):].lstrip('\n')
    if cleaned.startswith('```'):
        cleaned = cleaned[len('```'):].lstrip('\n')
    if cleaned.endswith('```'):
        cleaned = cleaned[:-3].rstrip('\n')

    # Check if input is likely Markdown (contains at least one heading, bullet, or bold)
    if not any(x in cleaned for x in ['##', '###', '-', '**']):
        # Log a warning and save a minimal DOCX with a warning message
        from docx import Document
        doc = Document()
        doc.add_paragraph('Warning: Resume content was not in Markdown format. Please regenerate.')
        doc.save(resume_file_path)
        print(f"[WARN] Resume content did not appear to be Markdown. Saved warning to {resume_file_path}")
        return resume_file_path

    # Convert Markdown to DOCX and save
    markdown_to_docx(cleaned, resume_file_path)
    return resume_file_path


# --- Resume-specific Markdown to DOCX conversion ---
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def markdown_to_docx(markdown_content, docx_path):
    """
    Converts Markdown resume content to a DOCX file with ATS-friendly formatting.
    - ## Header → bold, font size 14
    - ### Subheader → bold, font size 12
    - - item → bullet list
    - **bold** → inline bold text
    - No tables, no columns
    """
    doc = Document()
    # Set narrow margins to fit more content per page
    section = doc.sections[0]
    from docx.shared import Inches
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    bullet_mode = False
    for line in markdown_content.splitlines():
        line = line.rstrip()
        if not line.strip():
            bullet_mode = False
            continue
        # Section Header (##)
        if line.startswith('## '):
            header_text = line[3:].strip()
            # Remove all leading/trailing asterisks and spaces (e.g., '**Software Engineer**', '*Software Engineer*')
            header_text = header_text.strip('*').strip()
            para = doc.add_paragraph()
            run = para.add_run(header_text)
            run.bold = True
            run.font.size = Pt(12)  # Slightly smaller than before
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            bullet_mode = False
            continue
        # Subheader (###)
        if line.startswith('### '):
            subheader_text = line[4:].strip()
            subheader_text = subheader_text.strip('*').strip()
            para = doc.add_paragraph()
            run = para.add_run(subheader_text)
            run.bold = True
            run.font.size = Pt(11)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.space_before = Pt(1)
            bullet_mode = False
            continue
        # Bullet point
        if line.startswith('- '):
            para = doc.add_paragraph(style='List Bullet')
            _add_markdown_run(para, line[2:].strip(), font_size=10)
            para.paragraph_format.space_after = Pt(0.5)
            para.paragraph_format.space_before = Pt(0.5)
            bullet_mode = True
            continue
        # Normal paragraph (with possible bold)
        para = doc.add_paragraph()
        _add_markdown_run(para, line.strip(), font_size=10)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.space_before = Pt(1)
        bullet_mode = False
    # Support both file path and BytesIO for docx_path
    if hasattr(docx_path, 'write'):
        doc.save(docx_path)
    else:
        doc.save(str(docx_path))

def _add_markdown_run(para, text, font_size=10):
    """
    Helper to add text to a paragraph, handling **bold** inline markdown. Allows font size adjustment.
    """
    pattern = r'(\*\*[^*]+\*\*)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(font_size)
        else:
            run = para.add_run(part)
            run.font.size = Pt(font_size)
